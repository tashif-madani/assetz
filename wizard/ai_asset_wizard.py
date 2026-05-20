import base64
import io
import json
import logging

from odoo import api, fields, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    _logger.warning("PyPDF2 not installed. PDF parsing will be disabled.")

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    _logger.info("pdfplumber not installed. Advanced PDF parsing not available.")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    _logger.warning("python-docx not installed. Word document parsing will be disabled.")

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    _logger.warning("openpyxl not installed. Excel parsing will be disabled.")

# OCR dependencies for scanned PDFs
try:
    import pdf2image
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False
    _logger.info("pdf2image not installed. OCR for scanned PDFs not available.")

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    _logger.info("pytesseract not installed. OCR for scanned PDFs not available.")


class AIAssetWizard(models.TransientModel):
    """Wizard for creating assets from documents using AI extraction."""

    _name = "assetz.ai.asset.wizard"
    _description = "AI Asset Creation Wizard"

    # State
    state = fields.Selection(
        selection=[
            ("upload", "Upload Document"),
            ("preview", "Preview Extracted Data"),
            ("done", "Complete"),
        ],
        string="State",
        default="upload",
    )

    # Document upload
    document_data = fields.Binary(
        string="Document",
        help="Upload a PDF, Word, or Excel document containing asset information",
    )
    document_name = fields.Char(string="Document Name")
    document_type = fields.Selection(
        selection=[
            ("pdf", "PDF"),
            ("docx", "Word Document"),
            ("xlsx", "Excel Spreadsheet"),
        ],
        string="Document Type",
        compute="_compute_document_type",
        store=True,
    )

    # Extracted data
    extracted_text = fields.Text(
        string="Extracted Text",
        readonly=True,
    )

    # Asset lines (preview)
    asset_line_ids = fields.One2many(
        comodel_name="assetz.ai.asset.wizard.line",
        inverse_name="wizard_id",
        string="Extracted Assets",
    )

    # Document preview - same as document_data but used for preview display
    document_preview = fields.Binary(
        string="Document Preview",
        related="document_data",
        readonly=True,
    )

    # Results
    created_asset_ids = fields.Many2many(
        comodel_name="assetz.asset",
        relation="assetz_ai_wizard_created_assets_rel",
        string="Created Assets",
    )
    error_message = fields.Text(string="Error Message")

    @api.depends("document_name")
    def _compute_document_type(self):
        """Determine document type from file extension."""
        for wizard in self:
            if wizard.document_name:
                name_lower = wizard.document_name.lower()
                if name_lower.endswith(".pdf"):
                    wizard.document_type = "pdf"
                elif name_lower.endswith((".docx", ".doc")):
                    wizard.document_type = "docx"
                elif name_lower.endswith((".xlsx", ".xls")):
                    wizard.document_type = "xlsx"
                else:
                    wizard.document_type = False
            else:
                wizard.document_type = False

    @api.model
    def action_open_wizard(self):
        """Open the wizard from server action or menu."""
        return {
            "type": "ir.actions.act_window",
            "name": "Create Assets from Document (AI)",
            "res_model": "assetz.ai.asset.wizard",
            "view_mode": "form",
            "target": "new",
            "context": self.env.context,
        }

    def action_process_document(self):
        """Parse document and call AI to extract asset data."""
        self.ensure_one()

        if not self.document_data:
            raise UserError("Please upload a document first.")

        if not self.document_type:
            raise UserError(
                "Unsupported file type. Please upload a PDF, Word (.docx), "
                "or Excel (.xlsx) file."
            )

        # Check if AI is enabled
        ai_enabled = self.env["ir.config_parameter"].sudo().get_param(
            "assetz.enable_ai_recommendations", "False"
        )
        if ai_enabled != "True":
            raise UserError(
                "AI Recommendations are not enabled. Please enable them in "
                "Settings > Assetz > AI Recommendations."
            )

        # Parse document
        try:
            self.extracted_text = self._parse_document()
        except Exception as e:
            raise UserError(f"Failed to parse document: {str(e)}")

        if not self.extracted_text or len(self.extracted_text.strip()) < 50:
            raise UserError(
                "Could not extract sufficient text from the document. "
                "Please ensure the document contains readable text."
            )

        # Call AI for extraction
        try:
            extracted_assets = self._call_ai_extraction()
        except Exception as e:
            _logger.error(f"AI extraction failed: {e}")
            raise UserError(f"AI extraction failed: {str(e)}")

        # Clear existing lines
        self.asset_line_ids.unlink()

        # Create asset lines from AI response
        if extracted_assets and "assets" in extracted_assets:
            for idx, asset_data in enumerate(extracted_assets["assets"]):
                self._create_asset_line(asset_data, idx + 1)

        if not self.asset_line_ids:
            raise UserError(
                "AI could not extract any asset information from this document. "
                "Please ensure the document contains asset details."
            )

        self.state = "preview"

        return {
            "type": "ir.actions.act_window",
            "res_model": self._name,
            "res_id": self.id,
            "view_mode": "form",
            "target": "new",
        }

    def _parse_document(self):
        """Parse document and extract text based on type."""
        content = base64.b64decode(self.document_data)

        if self.document_type == "pdf":
            return self._parse_pdf(content)
        elif self.document_type == "docx":
            return self._parse_docx(content)
        elif self.document_type == "xlsx":
            return self._parse_xlsx(content)
        else:
            raise UserError("Unsupported document type.")

    def _parse_pdf(self, content):
        """Extract text from PDF file using multiple methods."""
        text_parts = []
        pdf_bytes = io.BytesIO(content)

        # Try pdfplumber first (better for complex layouts, tables, certificates)
        if HAS_PDFPLUMBER:
            try:
                _logger.info("Attempting PDF extraction with pdfplumber...")
                pdf_bytes.seek(0)
                with pdfplumber.open(pdf_bytes) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            # Extract text from page
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(f"--- Page {page_num + 1} ---")
                                text_parts.append(page_text)

                            # Also extract text from tables (common in certificates)
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    if row:
                                        row_text = " | ".join(
                                            str(cell).strip() for cell in row
                                            if cell and str(cell).strip()
                                        )
                                        if row_text:
                                            text_parts.append(row_text)
                        except Exception as e:
                            _logger.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                            continue

                extracted_text = "\n".join(text_parts)
                _logger.info(f"pdfplumber extraction: {len(extracted_text)} chars from {len(pdf.pages)} pages")

                if len(extracted_text.strip()) > 50:
                    return extracted_text
                else:
                    _logger.info("pdfplumber returned insufficient text, trying PyPDF2...")
                    text_parts = []  # Reset for next attempt
            except Exception as e:
                _logger.warning(f"pdfplumber extraction failed: {e}")
                text_parts = []

        # Fallback to PyPDF2
        if not HAS_PYPDF2:
            raise UserError(
                "PDF parsing libraries not available. "
                "Please install with: pip install pdfplumber PyPDF2"
            )

        try:
            _logger.info("Attempting PDF extraction with PyPDF2...")
            pdf_bytes.seek(0)
            reader = PdfReader(pdf_bytes)

            for page_num, page in enumerate(reader.pages):
                try:
                    # Try standard text extraction
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---")
                        text_parts.append(page_text)

                    # Try to extract text from annotations (form fields)
                    if "/Annots" in page:
                        try:
                            annots = page["/Annots"]
                            if annots:
                                for annot in annots:
                                    annot_obj = annot.get_object() if hasattr(annot, 'get_object') else annot
                                    if "/V" in annot_obj:
                                        value = annot_obj["/V"]
                                        if value and str(value).strip():
                                            text_parts.append(f"Form field: {value}")
                        except Exception:
                            pass

                except Exception as e:
                    _logger.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            extracted_text = "\n".join(text_parts)
            _logger.info(f"PyPDF2 extraction: {len(extracted_text)} chars from {len(reader.pages)} pages")

        except Exception as e:
            _logger.error(f"PyPDF2 extraction failed: {e}")
            extracted_text = ""

        if not extracted_text.strip():
            _logger.warning("No text extracted from PDF - attempting OCR for scanned document...")
            # Try OCR as last resort for scanned PDFs
            extracted_text = self._parse_pdf_with_ocr(content)

        return extracted_text

    def _parse_pdf_with_ocr(self, content):
        """Extract text from scanned PDF using OCR."""
        if not HAS_PDF2IMAGE or not HAS_TESSERACT:
            _logger.warning(
                "OCR libraries not available. Install with: "
                "pip install pdf2image pytesseract (and install Tesseract OCR)"
            )
            return ""

        text_parts = []
        try:
            _logger.info("Attempting OCR extraction for scanned PDF...")

            # Convert PDF pages to images
            images = pdf2image.convert_from_bytes(content, dpi=200)
            _logger.info(f"Converted PDF to {len(images)} images for OCR")

            for page_num, image in enumerate(images):
                try:
                    # Use Tesseract OCR to extract text from image
                    page_text = pytesseract.image_to_string(image)
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} (OCR) ---")
                        text_parts.append(page_text.strip())
                except Exception as e:
                    _logger.warning(f"OCR failed on page {page_num + 1}: {e}")
                    continue

            extracted_text = "\n".join(text_parts)
            _logger.info(f"OCR extraction: {len(extracted_text)} chars from {len(images)} pages")
            return extracted_text

        except Exception as e:
            _logger.error(f"OCR extraction failed: {e}")
            return ""

    def _parse_docx(self, content):
        """Extract text from Word document."""
        if not HAS_DOCX:
            raise UserError(
                "python-docx library is required for Word document parsing. "
                "Please install with: pip install python-docx"
            )

        doc = DocxDocument(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _parse_xlsx(self, content):
        """Extract text from Excel file."""
        if not HAS_OPENPYXL:
            raise UserError(
                "openpyxl library is required for Excel parsing. "
                "Please install with: pip install openpyxl"
            )

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        text_parts = []

        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None and str(cell).strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _call_ai_extraction(self):
        """Call AI provider to extract asset data from document text."""
        ai_provider = self.env["ir.config_parameter"].sudo().get_param(
            "assetz.ai_provider", "none"
        )

        if ai_provider == "openai":
            return self._call_openai_extraction()
        elif ai_provider == "anthropic":
            return self._call_anthropic_extraction()
        else:
            raise UserError(
                "No AI provider configured. Please select OpenAI or Anthropic "
                "in Settings > Assetz > AI Recommendations."
            )

    def _get_extraction_prompt(self):
        """Build the extraction prompt for AI."""
        # Get attribute level setting
        level = self.env["ir.config_parameter"].sudo().get_param(
            "assetz.ai_attribute_level", "2"
        )

        level_descriptions = {
            "1": "5 critical",
            "2": "10 standard",
            "3": "20 comprehensive",
        }
        max_attributes = level_descriptions.get(level, "10 standard")

        # Truncate text if too long
        doc_text = self.extracted_text
        if len(doc_text) > 15000:
            doc_text = doc_text[:15000] + "\n... [truncated]"

        prompt = f"""Analyze this document and extract information about physical assets
(equipment, machinery, IT hardware, vehicles, furniture, etc.).

Document Content:
---
{doc_text}
---

Extract ALL assets found in this document. Respond ONLY with valid JSON:

{{
    "assets": [
        {{
            "name": "Asset name (required)",
            "description": "Brief description or null",
            "serial_number": "Serial/model number or null",
            "model": "Model name/number or null",
            "manufacturer": "Manufacturer/brand or null",
            "acquisition_date": "YYYY-MM-DD or null",
            "acquisition_cost": numeric or null,
            "warranty_expiry_date": "YYYY-MM-DD or null",
            "useful_life_years": integer or null,
            "suggested_category": "Category name (e.g., Laptop, Server, Furniture)",
            "attributes": [
                {{
                    "name": "Attribute name",
                    "type": "char|integer|float|date|selection|boolean",
                    "value": "Extracted value",
                    "is_critical": true or false,
                    "options": ["opt1", "opt2"]
                }}
            ]
        }}
    ]
}}

Attribute Level: Maximum {max_attributes} attributes per asset.
Prioritize the most important/critical attributes first.
For "type", use: char (text), integer (whole number), float (decimal),
date (YYYY-MM-DD), selection (dropdown with options), boolean (true/false).
Only include "options" array for selection type.
Use null for unknown values. Ensure valid JSON output only."""

        return prompt

    def _call_openai_extraction(self):
        """Call OpenAI to extract asset data using GPT-4o model."""
        try:
            import openai
        except ImportError:
            raise UserError(
                "OpenAI library not installed. "
                "Please install with: pip install openai"
            )

        api_key = self.env["ir.config_parameter"].sudo().get_param(
            "assetz.openai_api_key"
        )
        if not api_key:
            raise UserError("OpenAI API key not configured in settings.")

        prompt = self._get_extraction_prompt()

        try:
            client = openai.OpenAI(api_key=api_key)

            # Use GPT-4o with structured output
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at extracting structured data from documents. "
                                   "Analyze documents thoroughly to identify all assets, their details, "
                                   "and relevant attributes. Always respond with valid JSON only.",
                    },
                    {"role": "user", "content": prompt},
                ],
                response_format={
                    "type": "json_schema",
                    "json_schema": {
                        "name": "asset_extraction",
                        "strict": True,
                        "schema": {
                            "type": "object",
                            "properties": {
                                "assets": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "name": {"type": "string"},
                                            "description": {"type": ["string", "null"]},
                                            "serial_number": {"type": ["string", "null"]},
                                            "model": {"type": ["string", "null"]},
                                            "manufacturer": {"type": ["string", "null"]},
                                            "acquisition_date": {"type": ["string", "null"]},
                                            "acquisition_cost": {"type": ["number", "null"]},
                                            "warranty_expiry_date": {"type": ["string", "null"]},
                                            "useful_life_years": {"type": ["integer", "null"]},
                                            "suggested_category": {"type": "string"},
                                            "attributes": {
                                                "type": "array",
                                                "items": {
                                                    "type": "object",
                                                    "properties": {
                                                        "name": {"type": "string"},
                                                        "type": {"type": "string"},
                                                        "value": {"type": ["string", "null"]},
                                                        "is_critical": {"type": "boolean"},
                                                        "options": {"type": "array", "items": {"type": "string"}}
                                                    },
                                                    "required": ["name", "type", "value", "is_critical", "options"],
                                                    "additionalProperties": False
                                                }
                                            }
                                        },
                                        "required": ["name", "description", "serial_number", "model", "manufacturer", "acquisition_date", "acquisition_cost", "warranty_expiry_date", "useful_life_years", "suggested_category", "attributes"],
                                        "additionalProperties": False
                                    }
                                }
                            },
                            "required": ["assets"],
                            "additionalProperties": False
                        }
                    }
                },
                temperature=0.1,
                max_tokens=4000,
            )

            result_text = response.choices[0].message.content.strip()

            if not result_text:
                raise UserError("No response received from AI model.")

            _logger.info(f"OpenAI GPT-4o response received: {len(result_text)} chars")

            return json.loads(result_text.strip())

        except json.JSONDecodeError as e:
            _logger.error(f"Failed to parse AI response as JSON: {e}")
            raise UserError(
                "AI returned invalid response. Please try again or check the document format."
            )
        except Exception as e:
            _logger.error(f"OpenAI API call failed: {e}")
            raise UserError(f"OpenAI API error: {str(e)}")

    def _call_anthropic_extraction(self):
        """Call Anthropic Claude to extract asset data."""
        try:
            import anthropic
        except ImportError:
            raise UserError(
                "Anthropic library not installed. "
                "Please install with: pip install anthropic"
            )

        api_key = self.env["ir.config_parameter"].sudo().get_param(
            "assetz.anthropic_api_key"
        )
        if not api_key:
            raise UserError("Anthropic API key not configured in settings.")

        prompt = self._get_extraction_prompt()

        try:
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=4000,
                messages=[
                    {"role": "user", "content": prompt},
                ],
            )

            result_text = response.content[0].text.strip()

            # Clean up potential markdown formatting
            if result_text.startswith("```"):
                result_text = result_text.split("```")[1]
                if result_text.startswith("json"):
                    result_text = result_text[4:]
            if result_text.endswith("```"):
                result_text = result_text[:-3]

            return json.loads(result_text.strip())

        except json.JSONDecodeError as e:
            _logger.error(f"Failed to parse AI response as JSON: {e}")
            raise UserError(
                "AI returned invalid response. Please try again or check the document format."
            )
        except Exception as e:
            _logger.error(f"Anthropic API call failed: {e}")
            raise UserError(f"Anthropic API error: {str(e)}")

    def _create_asset_line(self, asset_data, sequence):
        """Create an asset preview line from extracted data."""
        # Find or suggest category
        category_name = asset_data.get("suggested_category", "")
        category = False
        if category_name:
            category = self.env["assetz.category"].search(
                [("name", "=ilike", category_name)], limit=1
            )

        # Parse dates
        def parse_date(date_str):
            if not date_str:
                return False
            try:
                from datetime import datetime
                return datetime.strptime(date_str, "%Y-%m-%d").date()
            except (ValueError, TypeError):
                return False

        # Create asset line
        line_vals = {
            "wizard_id": self.id,
            "sequence": sequence,
            "include": True,
            "ai_name": asset_data.get("name", ""),
            "ai_description": asset_data.get("description") or "",
            "ai_serial_number": asset_data.get("serial_number") or "",
            "ai_model": asset_data.get("model") or "",
            "ai_manufacturer": asset_data.get("manufacturer") or "",
            "ai_acquisition_date": parse_date(asset_data.get("acquisition_date")),
            "ai_acquisition_cost": asset_data.get("acquisition_cost") or 0.0,
            "ai_warranty_expiry_date": parse_date(asset_data.get("warranty_expiry_date")),
            "ai_useful_life_years": asset_data.get("useful_life_years") or 0,
            "category_id": category.id if category else False,
            "ai_suggested_category": category_name,
            "create_new_category": bool(category_name and not category),
        }

        asset_line = self.env["assetz.ai.asset.wizard.line"].create(line_vals)

        # Create attribute lines
        attributes = asset_data.get("attributes", [])
        for attr in attributes:
            attr_type = attr.get("type", "char")
            # Map AI types to Odoo types
            type_map = {
                "char": "char",
                "text": "char",
                "string": "char",
                "integer": "integer",
                "int": "integer",
                "float": "float",
                "decimal": "float",
                "number": "float",
                "date": "date",
                "selection": "selection",
                "dropdown": "selection",
                "boolean": "boolean",
                "bool": "boolean",
            }
            attr_type = type_map.get(attr_type.lower(), "char")

            options = attr.get("options", [])
            option_str = ", ".join(str(o) for o in options) if options else ""

            self.env["assetz.ai.asset.wizard.attribute"].create({
                "asset_line_id": asset_line.id,
                "include": True,
                "name": attr.get("name", ""),
                "attribute_type": attr_type,
                "suggested_value": str(attr.get("value", "")) if attr.get("value") else "",
                "is_required": attr.get("is_critical", False),
                "option_values": option_str,
            })

        return asset_line

    def action_create_assets(self):
        """Create assets from the previewed data."""
        self.ensure_one()

        if not self.asset_line_ids.filtered("include"):
            raise UserError("Please select at least one asset to create.")

        created_assets = self.env["assetz.asset"]

        for line in self.asset_line_ids.filtered("include"):
            try:
                asset = line._create_asset()
                if asset:
                    created_assets |= asset

                    # Attach original document to asset
                    self.env["ir.attachment"].create({
                        "name": self.document_name or "Source Document",
                        "type": "binary",
                        "datas": self.document_data,
                        "res_model": "assetz.asset",
                        "res_id": asset.id,
                        "mimetype": self._get_mimetype(),
                    })
            except Exception as e:
                _logger.error(f"Failed to create asset from line {line.id}: {e}")
                self.error_message = (self.error_message or "") + f"\nError creating {line.ai_name}: {str(e)}"

        self.created_asset_ids = created_assets
        self.state = "done"

        return {
            "type": "ir.actions.act_window",
            "res_model": self._name,
            "res_id": self.id,
            "view_mode": "form",
            "target": "new",
        }

    def _get_mimetype(self):
        """Get MIME type for the uploaded document."""
        mimetypes = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
        return mimetypes.get(self.document_type, "application/octet-stream")

    def action_back(self):
        """Go back to upload state."""
        self.ensure_one()
        if self.state == "preview":
            self.state = "upload"
        return {
            "type": "ir.actions.act_window",
            "res_model": self._name,
            "res_id": self.id,
            "view_mode": "form",
            "target": "new",
        }

    def action_view_assets(self):
        """View the created assets."""
        self.ensure_one()
        if not self.created_asset_ids:
            return {"type": "ir.actions.act_window_close"}

        return {
            "type": "ir.actions.act_window",
            "name": "Created Assets",
            "res_model": "assetz.asset",
            "view_mode": "list,form",
            "domain": [("id", "in", self.created_asset_ids.ids)],
            "target": "current",
        }


class AIAssetWizardLine(models.TransientModel):
    """Preview line for an extracted asset."""

    _name = "assetz.ai.asset.wizard.line"
    _description = "AI Asset Wizard Line"
    _order = "sequence"

    wizard_id = fields.Many2one(
        comodel_name="assetz.ai.asset.wizard",
        string="Wizard",
        required=True,
        ondelete="cascade",
    )
    sequence = fields.Integer(string="Sequence", default=10)
    include = fields.Boolean(string="Include", default=True)

    # Asset fields (editable)
    ai_name = fields.Char(string="Asset Name", required=True)
    ai_description = fields.Text(string="Description")
    ai_serial_number = fields.Char(string="Serial Number")
    ai_model = fields.Char(string="Model")
    ai_manufacturer = fields.Char(string="Manufacturer")
    ai_acquisition_date = fields.Date(string="Acquisition Date")
    ai_acquisition_cost = fields.Float(string="Acquisition Cost")
    ai_warranty_expiry_date = fields.Date(string="Warranty Expiry")
    ai_useful_life_years = fields.Integer(string="Useful Life (Years)")
    location_id = fields.Many2one(
        comodel_name="assetz.location",
        string="Location",
    )

    # Category
    category_id = fields.Many2one(
        comodel_name="assetz.category",
        string="Category",
    )
    ai_suggested_category = fields.Char(
        string="Suggested Category",
        help="Category name suggested by AI. You can edit this before creating.",
    )
    create_new_category = fields.Boolean(
        string="Create New Category",
        help="Create a new category if it doesn't exist",
    )

    # Attributes
    attribute_line_ids = fields.One2many(
        comodel_name="assetz.ai.asset.wizard.attribute",
        inverse_name="asset_line_id",
        string="Attributes",
    )

    @api.onchange("category_id")
    def _onchange_category_id(self):
        """Update create_new_category based on category selection."""
        if self.category_id:
            self.create_new_category = False

    def action_view_details(self):
        """Open a dialog to view/edit the asset details and attributes."""
        self.ensure_one()
        return {
            "type": "ir.actions.act_window",
            "name": f"Asset Details: {self.ai_name}",
            "res_model": "assetz.ai.asset.wizard.line",
            "res_id": self.id,
            "view_mode": "form",
            "target": "new",
            "context": {"dialog_size": "large"},
        }

    def action_save_and_close(self):
        """Save changes and return to the parent wizard."""
        self.ensure_one()
        # Return action to reopen the parent wizard
        return {
            "type": "ir.actions.act_window",
            "res_model": "assetz.ai.asset.wizard",
            "res_id": self.wizard_id.id,
            "view_mode": "form",
            "target": "new",
        }

    def _create_asset(self):
        """Create an asset record from this line."""
        self.ensure_one()

        category = self.category_id

        # Handle category - check if it exists or needs to be created
        if self.create_new_category and self.ai_suggested_category:
            # Check if category with this name already exists
            existing_category = self.env["assetz.category"].search([
                ("name", "=ilike", self.ai_suggested_category.strip())
            ], limit=1)

            if existing_category:
                # Use existing category instead of creating new one
                category = existing_category
                _logger.info(f"Using existing category: {category.name}")
            else:
                # Create new category with auto-generated code
                category_name = self.ai_suggested_category.strip()
                # Generate code from first letters of each word (max 4 chars)
                words = category_name.split()
                if len(words) > 1:
                    code = "".join(w[0].upper() for w in words[:4])
                else:
                    code = category_name[:4].upper()

                # Ensure code is unique
                base_code = code
                counter = 1
                while self.env["assetz.category"].search([("code", "=", code)], limit=1):
                    code = f"{base_code}{counter}"
                    counter += 1

                category = self.env["assetz.category"].create({
                    "name": category_name,
                    "code": code,
                })
                _logger.info(f"Created new category: {category.name} with code: {code}")

        # Process attributes for the category
        if category:
            self._process_category_attributes(category)

        # Ensure we have a location (required field)
        location = self.location_id
        if not location:
            # Try to find a default location
            location = self.env["assetz.location"].search([], limit=1)
            if not location:
                # Create a default location
                location = self.env["assetz.location"].create({
                    "name": "Default Location",
                })

        # Create the asset
        asset_vals = {
            "name": self.ai_name,
            "description": self.ai_description,
            "serial_number": self.ai_serial_number,
            "model": self.ai_model,
            "manufacturer": self.ai_manufacturer,
            "acquisition_date": self.ai_acquisition_date,
            "acquisition_cost": self.ai_acquisition_cost,
            "warranty_expiry_date": self.ai_warranty_expiry_date,
            "useful_life_years": self.ai_useful_life_years,
            "category_id": category.id if category else False,
            "location_id": location.id,
        }

        asset = self.env["assetz.asset"].create(asset_vals)

        # Create attribute values if category exists
        if category:
            self._create_asset_attribute_values(asset, category)

        return asset

    def _process_category_attributes(self, category):
        """Process attributes - create new ones or add values to existing."""
        for attr_line in self.attribute_line_ids.filtered("include"):
            # Check if attribute already exists in category (case-insensitive)
            existing_attr = category.attribute_ids.filtered(
                lambda a: a.name.lower() == attr_line.name.lower()
            )

            if existing_attr:
                cat_attr = existing_attr[0]
                _logger.info(f"Attribute '{attr_line.name}' already exists in category")

                # If it's a selection type, check if we need to add new option
                if cat_attr.attribute_type == "selection" and attr_line.suggested_value:
                    self._ensure_attribute_option(cat_attr, attr_line.suggested_value)

                    # Also add any additional options
                    if attr_line.option_values:
                        for opt in attr_line.option_values.split(","):
                            opt_name = opt.strip()
                            if opt_name:
                                self._ensure_attribute_option(cat_attr, opt_name)
            else:
                # Create new attribute - always use selection type
                cat_attr = self.env["assetz.category.attribute"].create({
                    "category_id": category.id,
                    "name": attr_line.name,
                    "attribute_type": "selection",
                    "is_required": attr_line.is_required,
                })
                _logger.info(f"Created new attribute '{attr_line.name}' in category")

                # Collect all options: from option_values field + the suggested value
                options_set = set()

                # Add options from option_values field if any
                if attr_line.option_values:
                    for opt in attr_line.option_values.split(","):
                        opt_name = opt.strip()
                        if opt_name:
                            options_set.add(opt_name)

                # Add the suggested value as an option
                if attr_line.suggested_value:
                    options_set.add(attr_line.suggested_value.strip())

                # Create all options
                for opt_name in options_set:
                    self.env["assetz.category.attribute.option"].create({
                        "attribute_id": cat_attr.id,
                        "name": opt_name,
                    })

    def _ensure_attribute_option(self, cat_attr, option_value):
        """Ensure an option exists for a selection attribute, create if not."""
        if not option_value:
            return None

        option_value = option_value.strip()

        # Check if option already exists (case-insensitive)
        existing_option = cat_attr.option_ids.filtered(
            lambda o: o.name.lower() == option_value.lower()
        )

        if existing_option:
            return existing_option[0]

        # Create new option
        new_option = self.env["assetz.category.attribute.option"].create({
            "attribute_id": cat_attr.id,
            "name": option_value,
        })
        _logger.info(f"Added new option '{option_value}' to attribute '{cat_attr.name}'")
        return new_option

    def _create_asset_attribute_values(self, asset, category):
        """Create attribute values for the asset."""
        for attr_line in self.attribute_line_ids.filtered("include"):
            if not attr_line.suggested_value:
                continue

            # Find matching category attribute (case-insensitive)
            cat_attr = category.attribute_ids.filtered(
                lambda a: a.name.lower() == attr_line.name.lower()
            )

            if not cat_attr:
                continue

            cat_attr = cat_attr[0]
            attr_vals = {
                "asset_id": asset.id,
                "attribute_id": cat_attr.id,
                "value": attr_line.suggested_value,
            }

            # Handle selection type - find or create option
            if cat_attr.attribute_type == "selection":
                option = self._ensure_attribute_option(cat_attr, attr_line.suggested_value)
                if option:
                    attr_vals["value_selection_id"] = option.id
                    attr_vals["value"] = option.name

            self.env["assetz.asset.attribute.value"].create(attr_vals)


class AIAssetWizardAttribute(models.TransientModel):
    """Attribute line for an extracted asset."""

    _name = "assetz.ai.asset.wizard.attribute"
    _description = "AI Asset Wizard Attribute"
    _rec_name = "display_name"

    asset_line_id = fields.Many2one(
        comodel_name="assetz.ai.asset.wizard.line",
        string="Asset Line",
        required=True,
        ondelete="cascade",
    )
    include = fields.Boolean(string="Include", default=True)
    name = fields.Char(string="Attribute Name", required=True)
    attribute_type = fields.Selection(
        selection=[
            ("char", "Text"),
            ("integer", "Whole Number"),
            ("float", "Decimal Number"),
            ("date", "Date"),
            ("selection", "Dropdown"),
            ("boolean", "Yes/No"),
        ],
        string="Type",
        default="char",
    )
    suggested_value = fields.Char(string="Value")
    is_required = fields.Boolean(string="Required")
    option_values = fields.Char(
        string="Options",
        help="Comma-separated options for dropdown type",
    )
    display_name = fields.Char(
        string="Display Name",
        compute="_compute_display_name",
        store=False,
    )

    @api.depends("name", "suggested_value", "attribute_type")
    def _compute_display_name(self):
        """Compute display name for tags widget."""
        for attr in self:
            type_icons = {
                "char": "📝",
                "integer": "#️⃣",
                "float": "🔢",
                "date": "📅",
                "selection": "📋",
                "boolean": "✓",
            }
            icon = type_icons.get(attr.attribute_type, "")
            value_str = f" = {attr.suggested_value}" if attr.suggested_value else ""
            attr.display_name = f"{icon} {attr.name}{value_str}"
